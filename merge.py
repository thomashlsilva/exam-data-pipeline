from docx import Document
from docx.shared import Pt
import os


def merge_docx_files_with_titles(docx_folder, merged_docx_file):
    # Create a new Document to store the merged content
    merged_doc = Document()

    # Iterate through each file in the input folder
    for docx_file in sorted(os.listdir(docx_folder)):
        if docx_file.endswith(".docx"):
            # Add a page break before each new document
            if merged_doc.paragraphs:
                merged_doc.add_page_break()

            # Open the current DOCX file
            current_doc = Document(os.path.join(docx_folder, docx_file))

            # Add a title with the file name to the merged document
            title_paragraph = merged_doc.add_paragraph()
            title_run = title_paragraph.add_run(docx_file)
            title_run.bold = True
            title_run.font.size = Pt(14)  # You can adjust the font size if needed

            # Add each paragraph from the current document to the merged document
            for element in current_doc.element.body:
                merged_doc.element.body.append(element)

    # Save the merged document
    merged_doc.save(merged_docx_file)


if __name__ == "__main__":
    # Specify the folder containing the DOCX files and the path for the merged DOCX file
    docx_folder_path = "./docx"
    merged_docx_path = "merged_file.docx"

    merge_docx_files_with_titles(docx_folder_path, merged_docx_path)
